import streamlit as st
import os
import pandas as pd
from io import BytesIO
from docx import Document
from groq import Groq

# 1. Page & Layout Configuration
st.set_page_config(
    page_title="VocalKisan Dairy AI - DPR Assistant",
    page_icon="🥛",
    layout="wide"
)

# 2. Secure Secret Key Initialization
if "groq_key" not in st.session_state:
    if "GROQ_API_KEY" in st.secrets:
        st.session_state["groq_key"] = st.secrets["GROQ_API_KEY"]
    elif os.environ.get("GROQ_API_KEY"):
        st.session_state["groq_key"] = os.environ.get("GROQ_API_KEY")
    else:
        st.session_state["groq_key"] = ""

# Initialize local session state matrices for document compilation
if "dpr_chapters" not in st.session_state:
    st.session_state["dpr_chapters"] = {}

# 22-Chapter Master Index mapping directly to Institutional Framework
DPR_CHAPTERS_MAP = {
    "1. Executive Summary": "Executive summary outlining milestones, expected outcomes, and strategic interventions.",
    "2. Introduction": "Contextual introduction to the dairy project scope.",
    "3. Indian Dairy Sector Overview": "Overview of the sector, production systems, and current smallholder dynamics.",
    "4. Project Background & Need Assessment": "Problem statements regarding low productivity, feed deficits, and market channels.",
    "5. Project Objectives": "Targeted strategic objectives focused on sustainable productivity gains.",
    "6. Project Area & Beneficiary Analysis": "Geospatial assessment and demographic analysis of participating farmers.",
    "7. Baseline Assessment": "Existing status of animal assets, chilling facilities, and local cooperative infrastructure.",
    "8. Project Components": "Genetic Improvement, Ration Balancing (RBP), Village Procurement (VDCS), and Milk Quality Systems.",
    "9. Technical Design": "Operational workflows, AI doorstep service networks, and bulk milk cooling engineering plans.",
    "10. Infrastructure Requirements": "Civil works, plant and machinery parameters, laboratory testing kit matrices.",
    "11. Implementation Strategy": "Phased rollout schedules from pilot expansion to absolute optimization.",
    "12. Institutional Framework": "Governance backbone mapping central agencies, state federations, and district milk unions.",
    "13. Environmental & Social Management": "Waste tracking, effluent treatment optimization, and gender equity incentives.",
    "14. Financial Analysis": "CAPEX structure, recurring OPEX parameters, and means of finance models.",
    "15. Economic Analysis": "Societal returns, internal rate of return metrics, and benefit-cost balances.",
    "16. Procurement Strategy": "Package structures (PKG-1 to PKG-6), tendering processes, and L1/QCBS evaluation.",
    "17. Risk Analysis & Mitigation": "Technical data gaps, single-bid tracking, and feed volatility countermeasures.",
    "18. Monitoring & Evaluation Framework": "KPI dashboards tracking physical progress, financial burn rates, and operational yields.",
    "19. Project Implementation Schedule": "12-to-36 month milestone tracks and critical path indicators.",
    "20. Expected Outcomes": "Long-term production efficiency adjustments and domestic supply metrics up to 2050.",
    "21. Conclusions": "Investment viability final determinations and validation declarations.",
    "22. Annexures": "Detailed reference schedules, engineering single-line blueprints, and accounting balances."
}

# 3. Application Interface Layout
st.title("🌱 VocalKisan Dairy AI")
st.subheader("Automated Investment-Grade Detailed Project Report (DPR) Engine")
st.caption("Aligned with institutional structures (DIDF/NDDB standards) for long-term milk supply optimization.")

st.markdown("---")

# Left Control Dashboard Panel
with st.sidebar:
    st.header("⚙️ Configuration Controls")
    
    # Required Project Parameters
    project_location = st.text_input("Target Location / Region:", placeholder="e.g., Anand, Gujarat / Western Region")
    project_budget = st.number_input("Total Project Budget (in Lakhs INR):", min_value=5.0, value=50.0, step=5.0)
    project_timeframe = st.slider("Project Execution Horizon (Months):", min_value=6, max_value=60, value=12)
    
    focal_domain = st.selectbox(
        "Primary Focal Domain Focus:",
        options=[
            "Milk supply estimation by 2050 and self-sufficiency",
            "Roadmap to dairy productivity enhancement",
            "Milk and dairy products demand estimation",
            "Dairy processing infrastructure upgrading"
        ]
    )

    st.markdown("---")
    st.header("🔑 Authentication Status")
    
    user_key = st.text_input("Enter Groq API Key:", value=st.session_state["groq_key"], type="password")
    if user_key:
        st.session_state["groq_key"] = user_key
        st.success("Groq API credentials configured.")
    else:
        st.warning("Provide a key above or add it to secrets to execute generation.")
            
    st.markdown("---")
    st.caption("Powered by Open Source Python & Groq Llama-3 Pipelines.")

# Right Content Canvas Panel
col_main, col_export = st.columns([3, 1])

with col_main:
    st.header("📊 Interactive Compilation Desk")
    st.info("Select a chapter sequence block below to formulate comprehensive investment-grade text matrices dynamically.")
    
    selected_chapter = st.selectbox("Choose Target DPR Chapter to Draft:", options=list(DPR_CHAPTERS_MAP.keys()))
    st.markdown(f"**Focus Objective:** {DPR_CHAPTERS_MAP[selected_chapter]}")

    # Core AI Logic Generator Function using Groq Backend
    if st.button(f"Generate Content for {selected_chapter}"):
        if not st.session_state["groq_key"]:
            st.error("Error: Missing Groq API access token credentials.")
        else:
            with st.spinner(f"Computing deep data analysis matrix for {selected_chapter}..."):
                try:
                    # Instantiating the Groq Client
                    client = Groq(api_key=st.session_state["groq_key"])
                    
                    # Construct specific structural prompt context
                    prompt_message = f"""
                    You are an expert Senior Institutional Dairy Project Consultant and Agronomist.
                    Draft a professional, comprehensive, and investment-grade text for '{selected_chapter}' of a Detailed Project Report (DPR).
                    
                    Use these project configuration metrics explicitly in your generation parameters:
                    - Region/Location: {project_location}
                    - Total Budget Allocation: Rs. {project_budget} Lakhs INR
                    - Delivery Implementation Timeframe: {project_timeframe} Months
                    - Strategic Theme Vector: {focal_domain}
                    
                    Contextual Chapter Goal: {DPR_CHAPTERS_MAP[selected_chapter]}
                    
                    Ensure the writing is highly technical, authoritative, formal, and rich in structural depth. Avoid generic summaries or vague placeholders. If financial or technical metrics are described, generate plausible realistic inline data matrices matching institutional expectations (NDDB style).
                    """

                    # Groq Chat Completion Endpoint
                    response = client.chat.completions.create(
                        model="llama-3.3-70b-specdec",  # Ultra-fast, highly accurate model
                        messages=[
                            {"role": "system", "content": "You are a professional compiler of official institutional detailed project reports for agricultural infrastructure projects."},
                            {"role": "user", "content": prompt_message}
                        ],
                        temperature=0.2
                    )
                    
                    generated_text = response.choices[0].message.content
                    # Commit result vector directly into session cache
                    st.session_state["dpr_chapters"][selected_chapter] = generated_text
                    st.rerun()
                    
                except Exception as e:
                    st.error(f"Execution pipeline interrupted: {str(e)}")

    # Render current state of active chapter container
    if selected_chapter in st.session_state["dpr_chapters"]:
        st.markdown("---")
        st.subheader("🔍 Active Draft Preview Window")
        st.markdown(st.session_state["dpr_chapters"][selected_chapter])
    else:
        st.caption("No draft generated for this segment yet. Press the action button above to initialize.")

# Export Engine Control Panel
with col_export:
    st.header("💾 Document Export")
    st.markdown("Compile all generated structural segments into an integrated Word document file.")
    
    # Calculate current progress metric
    completed_count = len(st.session_state["dpr_chapters"])
    st.metric(label="Chapters Drafted", value=f"{completed_count} / 22")
    
    if completed_count > 0:
        # Instantiating docx buffer flow
        doc = Document()
        doc.add_heading(f"Detailed Project Report: {focal_domain}", level=0)
        doc.add_paragraph(f"Geospatial Context: {project_location}")
        doc.add_paragraph(f"Financial Footprint: Rs. {project_budget} Lakhs | Operational Window: {project_timeframe} Months")
        doc.add_paragraph("Generated via VocalKisan Dairy AI Open-Source Platform Architecture.")
        
        # Sequentially attach sorted session state components
        for ch_title in sorted(DPR_CHAPTERS_MAP.keys(), key=lambda x: int(x.split(".")[0])):
            if ch_title in st.session_state["dpr_chapters"]:
                doc.add_heading(ch_title, level=1)
                doc.add_paragraph(st.session_state["dpr_chapters"][ch_title])
                doc.add_page_break()
        
        # Save to memory stream block for clean client download action
        bio = BytesIO()
        doc.save(bio)
        bio.seek(0)
        
        st.download_button(
            label="Download Complete Word File (.docx)",
            data=bio,
            file_name=f"VocalKisan_Dairy_AI_DPR_{project_location.replace(' ', '_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        st.warning("Generate at least one individual chapter block above to unlock the primary Word document compilation download.")
